#!/usr/bin/env python3
"""Genera el documento Word de documentación del pipeline RETO."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
from pathlib import Path

OUTPUT = Path(__file__).parent / "Documentacion_Pipeline_RETO_2026-02-11.docx"


def set_cell(cell, text, bold=False, size=9):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, size=9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell(table.rows[r_idx + 1].cells[c_idx], str(val), size=9)
    return table


def build():
    doc = Document()

    # Estilos base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ============================================================
    # PORTADA
    # ============================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Proyecto RETO\nDocumentación Técnica del Pipeline de Datos")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Automatización, Base de Datos y Arquitectura")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = date_p.add_run(f"Fecha: 11 de febrero de 2026")
    run3.font.size = Pt(12)

    doc.add_page_break()

    # ============================================================
    # ÍNDICE
    # ============================================================
    doc.add_heading("Índice", level=1)
    toc_items = [
        "1. Resumen ejecutivo",
        "2. Fase 1: Automatización del pipeline (cron)",
        "3. Fase 2: Base de datos PostgreSQL",
        "4. Arquitectura general y fases futuras",
        "5. Archivos creados",
        "6. Configuración de cron",
        "7. Dependencias instaladas",
        "8. Anexo: esquema SQL completo",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    # ============================================================
    # 1. RESUMEN EJECUTIVO
    # ============================================================
    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "Se implementaron las dos primeras fases del sistema de datos del proyecto RETO, "
        "cuyo objetivo es automatizar la detección y análisis de discurso de odio en redes sociales "
        "(X/Twitter y YouTube)."
    )
    doc.add_paragraph(
        "Fase 1 — Automatización: Se creó un script maestro que ejecuta diariamente (a las 10:00 AM "
        "mediante cron) los 7 scripts del pipeline de datos de X en orden secuencial, con logging "
        "detallado y tolerancia a fallos (si un script falla, se registra el error y se continúa "
        "con el siguiente)."
    )
    doc.add_paragraph(
        "Fase 2 — Base de datos: Se diseñaron e implementaron 6 tablas en PostgreSQL (base reto_db) "
        "organizadas en dos schemas (raw y processed). Se creó un script de carga (paso 8 del pipeline) "
        "que lee los CSVs generados y los sube a PostgreSQL mediante UPSERT. Los CSVs siguen "
        "generándose como respaldo físico."
    )

    # ============================================================
    # 2. FASE 1: AUTOMATIZACIÓN
    # ============================================================
    doc.add_heading("2. Fase 1: Automatización del pipeline (cron)", level=1)

    doc.add_heading("2.1 Orden de ejecución", level=2)
    doc.add_paragraph(
        "El pipeline ejecuta los siguientes scripts en orden estricto. "
        "Ante fallo de cualquier script, se registra el error en el log y se continúa "
        "con el siguiente (opción B)."
    )

    scripts = [
        ("1", "sync_drive_csvs.py", "X_Mensajes/", "Descarga CSVs desde Google Drive (Apify)"),
        ("2", "consolidar_csv.py", "X_Mensajes/", "Consolida CSVs en un master unificado"),
        ("3", "filter_and_anonymize_x.py", "X_Mensajes/Anon/", "Anonimiza y detecta candidatos"),
        ("4", "X_terms_sheet.py", "Medios/", "Detecta términos de odio"),
        ("5", "score_baseline.py", "Etiquetado_Modelos/", "Scoring con modelo TF-IDF + LogReg"),
        ("6", "scored_prioridad_alta.py", "Etiquetado_Modelos/", "Filtra mensajes de prioridad alta"),
        ("7", "etiquetar_completo_llm.py", "Medios/ML/etiquetado_llm/", "Etiquetado con LLM (GPT)"),
        ("8", "load_to_db.py", "automatizacion_diaria/", "Carga CSVs a PostgreSQL"),
    ]
    add_table(doc, ["Paso", "Script", "Ubicación", "Descripción"], scripts)

    doc.add_heading("2.2 Script maestro", level=2)
    doc.add_paragraph(
        "Archivo: Clases/RETO/automatizacion_diaria/run_pipeline_diario.py"
    )
    doc.add_paragraph(
        "Funcionalidades:"
    )
    bullets = [
        "Ejecuta los 8 scripts en orden secuencial.",
        "Usa la variable de entorno PYTHON_BIN para indicar el intérprete Python (venv).",
        "Genera un log diario en automatizacion_diaria/logs/pipeline_YYYY-MM-DD.log.",
        "Timeout de 1 hora por script.",
        "Si un script falla, registra el error y continúa con el siguiente.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("2.3 Logs", level=2)
    doc.add_paragraph(
        "Cada ejecución genera dos archivos de log:"
    )
    log_rows = [
        ("pipeline_YYYY-MM-DD.log", "Detalle de cada script: inicio, fin, estado, errores"),
        ("cron_stdout.log", "Salida estándar acumulativa de todas las ejecuciones de cron"),
        ("load_db_YYYY-MM-DD.log", "Detalle de la carga a PostgreSQL por tabla"),
    ]
    add_table(doc, ["Archivo", "Contenido"], log_rows)

    # ============================================================
    # 3. FASE 2: BASE DE DATOS
    # ============================================================
    doc.add_heading("3. Fase 2: Base de datos PostgreSQL", level=1)

    doc.add_heading("3.1 Conexión", level=2)
    conn_rows = [
        ("Host", "localhost"),
        ("Puerto", "5432"),
        ("Base de datos", "reto_db"),
        ("Usuario", "postgres"),
        ("Credenciales", "Archivo .env en automatizacion_diaria/"),
    ]
    add_table(doc, ["Parámetro", "Valor"], conn_rows)

    doc.add_heading("3.2 Schemas", level=2)
    doc.add_paragraph(
        "Se crearon dos schemas nuevos sin modificar los existentes (delitos, public, reto):"
    )
    schema_rows = [
        ("raw", "Mensajes crudos tal como llegan del scraping, con datos de autor originales"),
        ("processed", "Mensajes anonimizados, scores de modelos, etiquetas LLM, validaciones manuales y resumen diario"),
        ("delitos", "(Existente) Estadísticas de delitos de odio - no modificado"),
        ("reto", "(Existente) Modelo de datos anterior - no modificado"),
    ]
    add_table(doc, ["Schema", "Descripción"], schema_rows)

    doc.add_heading("3.3 Tablas creadas", level=2)

    # raw.mensajes
    doc.add_heading("raw.mensajes", level=3)
    doc.add_paragraph("Mensajes crudos con datos de autor original. Fuente: consolidar_csv.py.")
    raw_cols = [
        ("message_uuid", "UUID", "PK", "Identificador único del mensaje"),
        ("platform", "VARCHAR(20)", "NOT NULL", "twitter, youtube"),
        ("tweet_id", "VARCHAR(50)", "", "ID original del tweet/comentario"),
        ("created_at", "TIMESTAMPTZ", "", "Fecha de publicación"),
        ("content_original", "TEXT", "NOT NULL", "Texto completo del mensaje"),
        ("author_username", "VARCHAR(100)", "", "Usuario original (dato sensible)"),
        ("author_id", "VARCHAR(50)", "", "ID original del autor"),
        ("source_media", "VARCHAR(200)", "", "Medio de comunicación fuente"),
        ("batch_id", "VARCHAR(100)", "", "Lote de scraping"),
        ("scrape_date", "TIMESTAMPTZ", "", "Fecha de extracción"),
        ("language", "VARCHAR(10)", "", "Idioma detectado"),
        ("url", "TEXT", "", "URL del mensaje original"),
        ("retweet_count", "INTEGER", "DEFAULT 0", ""),
        ("reply_count", "INTEGER", "DEFAULT 0", ""),
        ("like_count", "INTEGER", "DEFAULT 0", ""),
        ("quote_count", "INTEGER", "DEFAULT 0", ""),
        ("ingested_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Fecha de inserción en BD"),
    ]
    add_table(doc, ["Columna", "Tipo", "Restricción", "Descripción"], raw_cols)

    # processed.mensajes
    doc.add_heading("processed.mensajes", level=3)
    doc.add_paragraph("Mensajes anonimizados con detección de términos. Fuente: filter_and_anonymize_x.py. FK → raw.mensajes.")
    proc_cols = [
        ("message_uuid", "UUID", "PK, FK", "→ raw.mensajes"),
        ("platform", "VARCHAR(20)", "NOT NULL", ""),
        ("content_original", "TEXT", "NOT NULL", "Texto (sin datos de autor)"),
        ("source_media", "VARCHAR(200)", "", ""),
        ("created_at", "TIMESTAMPTZ", "", ""),
        ("language", "VARCHAR(10)", "", ""),
        ("url", "TEXT", "", ""),
        ("author_id_anon", "VARCHAR(64)", "", "Hash anónimo del ID"),
        ("author_username_anon", "VARCHAR(64)", "", "Hash anónimo del username"),
        ("matched_terms", "TEXT", "", "Términos de odio detectados"),
        ("has_hate_terms_match", "BOOLEAN", "DEFAULT FALSE", ""),
        ("match_count", "INTEGER", "DEFAULT 0", "Cantidad de términos encontrados"),
        ("strong_phrase", "BOOLEAN", "DEFAULT FALSE", "Frase fuerte detectada"),
        ("is_candidate", "BOOLEAN", "DEFAULT FALSE", "Es candidato a odio"),
        ("candidate_reason", "TEXT", "", "Motivo de candidatura"),
        ("processed_at", "TIMESTAMPTZ", "DEFAULT NOW()", ""),
    ]
    add_table(doc, ["Columna", "Tipo", "Restricción", "Descripción"], proc_cols)

    # processed.scores
    doc.add_heading("processed.scores", level=3)
    doc.add_paragraph("Resultados de modelos de scoring. PK compuesta permite múltiples modelos por mensaje. Fuente: score_baseline.py.")
    score_cols = [
        ("message_uuid", "UUID", "PK, FK", "→ processed.mensajes"),
        ("model_version", "VARCHAR(100)", "PK", "Ej: baseline_tfidf_logreg_v1"),
        ("proba_odio", "DOUBLE PRECISION", "", "Probabilidad de odio (0-1)"),
        ("pred_odio", "INTEGER", "", "Predicción binaria (0/1)"),
        ("priority", "VARCHAR(10)", "", "alta, media, baja"),
        ("score_date", "TIMESTAMPTZ", "DEFAULT NOW()", ""),
    ]
    add_table(doc, ["Columna", "Tipo", "Restricción", "Descripción"], score_cols)

    # processed.etiquetas_llm
    doc.add_heading("processed.etiquetas_llm", level=3)
    doc.add_paragraph("Clasificación del LLM. PK compuesta permite versionar el modelo/prompt. Fuente: etiquetar_completo_llm.py.")
    llm_cols = [
        ("message_uuid", "UUID", "PK, FK", "→ processed.mensajes"),
        ("llm_version", "VARCHAR(50)", "PK", "Versión del modelo/prompt (v1, v2...)"),
        ("clasificacion_principal", "VARCHAR(20)", "", "ODIO / NO_ODIO / DUDOSO"),
        ("categoria_odio_pred", "VARCHAR(100)", "", "Una de las 6 categorías ReTo"),
        ("intensidad_pred", "VARCHAR(5)", "", "1, 2 o 3"),
        ("resumen_motivo", "TEXT", "", "Explicación del LLM"),
        ("etiquetado_date", "TIMESTAMPTZ", "DEFAULT NOW()", ""),
    ]
    add_table(doc, ["Columna", "Tipo", "Restricción", "Descripción"], llm_cols)

    # processed.validaciones_manuales
    doc.add_heading("processed.validaciones_manuales", level=3)
    doc.add_paragraph(
        "Correcciones humanas del etiquetado LLM. Replica las columnas H-M del Google Sheet. "
        "Una validación por mensaje."
    )
    val_cols = [
        ("message_uuid", "UUID", "PK, FK", "→ processed.mensajes"),
        ("odio_flag", "BOOLEAN", "", "¿Es odio? (sí/no)"),
        ("categoria_odio", "VARCHAR(100)", "", "Categoría validada por el humano"),
        ("intensidad", "SMALLINT", "CHECK 1-3", "Intensidad validada"),
        ("humor_flag", "BOOLEAN", "", "¿Contiene humor? (sí/no)"),
        ("annotator_id", "VARCHAR(50)", "", "Quién validó"),
        ("annotation_date", "DATE", "", "Cuándo se validó"),
        ("coincide_con_llm", "BOOLEAN", "", "¿El LLM acertó?"),
        ("ingested_at", "TIMESTAMPTZ", "DEFAULT NOW()", ""),
    ]
    add_table(doc, ["Columna", "Tipo", "Restricción", "Descripción"], val_cols)

    doc.add_paragraph("Categorías válidas de odio (ReTo):")
    cats = [
        "odio_etnico_cultural_religioso",
        "odio_genero_identidad_orientacion",
        "odio_condicion_social_economica_salud",
        "odio_ideologico_politico",
        "odio_personal_generacional",
        "odio_profesiones_roles_publicos",
    ]
    for c in cats:
        doc.add_paragraph(c, style="List Bullet")

    # processed.resumen_diario
    doc.add_heading("processed.resumen_diario", level=3)
    doc.add_paragraph("Métricas agregadas por día y plataforma. Se recalcula al final de cada ejecución del pipeline.")
    res_cols = [
        ("fecha", "DATE", "PK", ""),
        ("platform", "VARCHAR(20)", "PK", ""),
        ("total_mensajes_raw", "INTEGER", "", ""),
        ("total_candidatos", "INTEGER", "", ""),
        ("total_odio_baseline", "INTEGER", "", ""),
        ("total_odio_llm", "INTEGER", "", ""),
        ("score_promedio", "DOUBLE PRECISION", "", ""),
    ]
    add_table(doc, ["Columna", "Tipo", "Restricción", "Descripción"], res_cols)

    doc.add_heading("3.4 Datos cargados (11/02/2026)", level=2)
    data_rows = [
        ("raw.mensajes", "29,908"),
        ("processed.mensajes", "29,908"),
        ("processed.scores", "17,813"),
        ("processed.etiquetas_llm", "2,489"),
        ("processed.resumen_diario", "1"),
        ("processed.validaciones_manuales", "0 (pendiente)"),
    ]
    add_table(doc, ["Tabla", "Filas"], data_rows)

    # ============================================================
    # 4. ARQUITECTURA
    # ============================================================
    doc.add_heading("4. Arquitectura general y fases futuras", level=1)

    doc.add_paragraph(
        "El sistema está diseñado en capas incrementales. Cada fase es funcional por sí misma "
        "y se puede usar en producción mientras se avanza a la siguiente."
    )

    fases = [
        ("Fase 1", "Completada", "Pipeline automatizado con cron (8 scripts, 10:00 AM diario)"),
        ("Fase 2", "Completada", "PostgreSQL local (schemas raw + processed, carga automática)"),
        ("Fase 3", "Pendiente", "Dashboard de visualización (Streamlit o Dash) conectado a PostgreSQL"),
        ("Fase 4", "Pendiente", "Migración a servidor Linux (cambio de conexión + pg_dump/restore)"),
        ("Fase 5", "Pendiente", "API FastAPI + Frontend React como sección de la web final"),
    ]
    add_table(doc, ["Fase", "Estado", "Descripción"], fases)

    doc.add_paragraph()
    doc.add_paragraph(
        "Flujo diario actual:\n"
        "1. Mac encendido a las 10:00 AM.\n"
        "2. Cron ejecuta run_pipeline_diario.py.\n"
        "3. Scripts 1-7 generan/actualizan CSVs (respaldo físico).\n"
        "4. Script 8 (load_to_db.py) lee CSVs y hace UPSERT a PostgreSQL.\n"
        "5. Logs quedan en automatizacion_diaria/logs/."
    )

    # ============================================================
    # 5. ARCHIVOS CREADOS
    # ============================================================
    doc.add_heading("5. Archivos creados", level=1)

    files = [
        ("run_pipeline_diario.py", "automatizacion_diaria/", "Script maestro del pipeline"),
        ("load_to_db.py", "automatizacion_diaria/", "Carga CSVs → PostgreSQL"),
        ("db_utils.py", "automatizacion_diaria/", "Conexión y funciones UPSERT genéricas"),
        (".env", "automatizacion_diaria/", "Credenciales PostgreSQL (no subir a git)"),
        ("schema_reto_db.sql", "automatizacion_diaria/", "Script SQL para recrear todas las tablas"),
        ("README.md", "automatizacion_diaria/", "Instrucciones de uso y configuración de cron"),
        ("logs/", "automatizacion_diaria/", "Carpeta de logs diarios"),
    ]
    add_table(doc, ["Archivo", "Ubicación", "Descripción"], files)

    # ============================================================
    # 6. CONFIGURACIÓN DE CRON
    # ============================================================
    doc.add_heading("6. Configuración de cron", level=1)

    doc.add_paragraph("Entrada actual en crontab del usuario:")
    doc.add_paragraph(
        '0 10 * * * cd "/Users/alejandroyankilevich/Documents/MASTER DATA SCIENCE" '
        '&& PYTHON_BIN="/Users/alejandroyankilevich/Documents/MASTER DATA SCIENCE/'
        'Clases/RETO/X_Mensajes/venv/bin/python3" /opt/homebrew/bin/python3 '
        '"Clases/RETO/automatizacion_diaria/run_pipeline_diario.py" >> '
        '"Clases/RETO/automatizacion_diaria/logs/cron_stdout.log" 2>&1',
        style="No Spacing"
    )

    doc.add_paragraph()
    doc.add_paragraph("Requisitos para que funcione:")
    reqs = [
        "Mac encendido y con sesión de usuario iniciada a las 10:00 AM.",
        "PostgreSQL corriendo en localhost:5432.",
        "El venv de X_Mensajes debe tener todas las dependencias instaladas.",
    ]
    for r in reqs:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph("Para editar la hora u otros parámetros: crontab -e")

    # ============================================================
    # 7. DEPENDENCIAS
    # ============================================================
    doc.add_heading("7. Dependencias instaladas", level=1)
    doc.add_paragraph(
        "Se instalaron las siguientes dependencias en el venv del proyecto "
        "(Clases/RETO/X_Mensajes/venv/):"
    )
    deps = [
        ("psycopg2-binary", "Conexión a PostgreSQL"),
        ("python-dotenv", "Lectura de archivos .env"),
        ("scikit-learn", "Modelo baseline (TF-IDF + LogReg)"),
        ("joblib", "Serialización de modelos"),
        ("python-docx", "Generación de este documento"),
    ]
    add_table(doc, ["Paquete", "Uso"], deps)

    # ============================================================
    # 8. ANEXO: SQL
    # ============================================================
    doc.add_heading("8. Anexo: esquema SQL completo", level=1)
    doc.add_paragraph(
        "El archivo schema_reto_db.sql contiene el DDL completo para recrear las tablas. "
        "Útil para la migración al servidor (Fase 4)."
    )
    doc.add_paragraph("Ubicación: Clases/RETO/automatizacion_diaria/schema_reto_db.sql")
    doc.add_paragraph("Uso:")
    doc.add_paragraph(
        "psql -h <host> -U postgres -d reto_db -f schema_reto_db.sql",
        style="No Spacing"
    )

    # ============================================================
    # CORRECCIONES APLICADAS
    # ============================================================
    doc.add_heading("Nota: correcciones aplicadas durante la sesión", level=2)
    fixes = [
        "filter_and_anonymize_x.py: se escapó un carácter '%' en un help de argparse "
        "que causaba error en Python 3.14 (línea 461: '(3%)' → '(3%%)').",
        "17 filas del master CSV tenían content_original vacío; se filtran automáticamente "
        "en load_to_db.py antes de insertar en PostgreSQL.",
    ]
    for f in fixes:
        doc.add_paragraph(f, style="List Bullet")

    # Guardar
    doc.save(str(OUTPUT))
    print(f"Documento generado: {OUTPUT}")


if __name__ == "__main__":
    build()
